from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

from app.common.enums import CVTemplate, JSONData, Number, Style, Table

from .common_utils import add_bullet_list, add_text
from .font_utils import set_raleway, set_raleway_medium
from .style_utils import add_style_table


def add_experience_left_cell(cell, experience):
    paragraph = cell.paragraphs[Number.ZERO.value]
    paragraph.add_run('\n')
    add_text(
        paragraph, f"{experience.get(JSONData.ROLE.value, '')}\n",
        set_raleway_medium, font_size=Style.NINE.value
    )
    add_text(
        paragraph, experience.get(JSONData.DATES.value, ''),
        set_raleway, font_size=Style.NINE.value
    )


def add_experience_right_cell(cell, experience, index):
    paragraph = cell.paragraphs[Number.ZERO.value]
    paragraph.add_run('\n')
    for title, key, new_line in JSONData.PROJECT.value[index]:
        add_text(
            paragraph, f'{title}', set_raleway_medium,
            font_size=Style.NINE.value
        )
        add_text(
            paragraph, experience.get(key, ''),
            set_raleway, font_size=Style.NINE.value, new_line=new_line
        )

    for title, key in JSONData.TASKS_ACHIEVEMENTS.value[index]:
        section_paragraph = cell.add_paragraph()
        add_text(
            section_paragraph, f'\n{title}', set_raleway_medium,
            font_size=Style.NINE.value
        )
        add_bullet_list(
            cell, experience.get(key, []), set_raleway,
            font_size=Style.NINE.value, bullet_color=Style.HEX_GREY.value
        )

    for title, key in JSONData.TEAM_STACK.value[index]:
        section_paragraph = cell.add_paragraph()
        add_text(
            section_paragraph, title, set_raleway_medium,
            font_size=Style.NINE.value
        )
        add_text(
            section_paragraph, experience.get(key, ''),
            set_raleway, font_size=Style.NINE.value
        )

    cell.add_paragraph()


def create_experiences_table(document, experiences, template_choice, index):
    table = document.add_table(rows=Table.ROWS.value, cols=Table.COLS.value)
    table.style = Table.STYLE.value

    for experience in experiences:
        row = table.add_row()
        left_cell = row.cells[Number.ZERO.value]
        right_cell = row.cells[Number.ONE.value]

        # Set column width
        left_cell.width = Inches(2.0)
        right_cell.width = Inches(4.0)

        add_experience_left_cell(left_cell, experience)
        add_experience_right_cell(right_cell, experience, index)

    if template_choice == CVTemplate.BUSINESSMATIKA.value:
        add_style_table(table, Table.BM_TOP.value, Table.BM_BOTTOM.value)
    elif template_choice == CVTemplate.TELESCOPE.value:
        add_style_table(table, Table.TELESCOPE_TOP.value, Table.TELESCOPE_BOTTOM.value)
    elif template_choice == CVTemplate.HUNTERCORE.value:
        add_style_table(table, Table.HUNTERCORE_TOP.value, Table.HUNTERCORE_BOTTOM.value)
    else:
        add_style_table(table, Table.BM_TOP.value, Table.BM_BOTTOM.value)


def create_skills_table(document, skills):
    """
    Creates table with skills, evenly distributed across columns
    (maximum 3 columns).
    """
    if not skills:
        return

    sorted_skills = sorted(skills)
    max_columns = Table.MAX_COLS.value
    rows = (len(sorted_skills) + max_columns - Number.ONE.value) // max_columns

    table = document.add_table(rows=rows, cols=max_columns)
    table.style = Table.STYLE.value

    # Заполняем таблицу
    skill_index = Number.ZERO.value
    for row in table.rows:
        for cell in row.cells:
            if skill_index < len(sorted_skills):
                skill = sorted_skills[skill_index]
                add_skill_to_cell(cell, skill)
                skill_index += Number.ONE.value
            else:
                cell.text = ''

    add_style_table(table, Table.TRANSCEND.value)


def add_skill_to_cell(cell, skill):
    """
    Adds skill to the given table cell with style and centering.
    """
    cell.text = ''
    paragraph = cell.paragraphs[Number.ZERO.value]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    skill_run = paragraph.add_run(skill)
    set_raleway(skill_run, font_size=Style.NINE.value)
