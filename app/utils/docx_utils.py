import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from app.common.enums import CVTemplate, CVTranslation, JSONData, Style
from app.config import (BM_FOOTER_PATH, BM_HEADER_PATH, HUNTERCORE_HEADER_PATH,
                        TELESCOPE_HEADER_PATH, logger)

from .common_utils import add_bullet_list, add_section, add_text
from .font_utils import set_raleway, set_raleway_medium
from .style_utils import add_images_to_header_footer
from .table_utils import create_experiences_table, create_skills_table


def generate_docx_from_json(data, output_stream, template_choice, language_choice):
    document = Document()
    if not template_choice:
        logger.error(
            f'Invalid template choice in generate_docx_from_json: "{template_choice}". '
            f'Default template "{CVTemplate.BUSINESSMATIKA.value} will be used.'
        )
        template_choice = CVTemplate.BUSINESSMATIKA.value

    if not language_choice:
        logger.error(
            f'Invalid language choice in generate_docx_from_json: "{language_choice}". '
            f'Default language "{CVTranslation.RUSSIAN.value} will be used.'
        )
        template_choice = CVTranslation.RUSSIAN.value

    index = 0 if language_choice == 'russian' else 1

    if template_choice == CVTemplate.BUSINESSMATIKA.value:
        add_images_to_header_footer(
            document,
            BM_HEADER_PATH, Style.BM_HEADER_WD.value, Style.BM_HEADER_H.value,
            BM_FOOTER_PATH, Style.BM_FOOTER_WD.value, Style.BM_FOOTER_H.value
        )
    elif template_choice == CVTemplate.HUNTERCORE.value:
        add_images_to_header_footer(
            document, HUNTERCORE_HEADER_PATH, Style.HUNT_HEADER_WD.value,
            Style.HUNT_HEADER_H.value
        )
    elif template_choice == CVTemplate.TELESCOPE.value:
        add_images_to_header_footer(
            document, TELESCOPE_HEADER_PATH, Style.TEL_HEADER_WD.value,
            Style.TEL_HEADER_H.value
        )
    
    header = data[JSONData.HEADER.value]
    sections = data[JSONData.SECTIONS.value]

    # Full name, Job title
    header_paragraph = document.add_paragraph()
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_text(
        header_paragraph, header.get(JSONData.FULL_NAME.value, ''),
        set_raleway_medium, font_size=Style.TWENTY.value
    )
    add_text(
        header_paragraph, f"\n{header.get(JSONData.JOB_TITLE.value, '')}",
        set_raleway, font_size=Style.TEN.value
    )

    # Grade
    add_section(
        document, JSONData.GRADE_TITLE.value[index],
        header.get(JSONData.GRADE_DATA.value, ''),
        set_raleway_medium, set_raleway
    )

    # Languages
    languages = sections[JSONData.LANGUAGES.value]
    add_section(
        document, languages[JSONData.TITLE.value],
        None, set_raleway_medium, set_raleway
    )
    add_bullet_list(
        document, languages.get(JSONData.ITEMS.value, []),
        set_raleway, font_size=Style.NINE.value,
        bullet_color=Style.HEX_GREY.value
    )

    # Skills
    skills = sections[JSONData.SKILLS.value]
    add_section(
        document, skills[JSONData.TITLE.value],
        None, set_raleway_medium, set_raleway
    )
    create_skills_table(document, skills.get(JSONData.ITEMS.value, []))

    # About, Repository
    add_section(
        document, JSONData.DESCR.value[index],
        header.get(JSONData.ABOUT.value, ''),
        set_raleway_medium, set_raleway
    )
    add_section(
        document, JSONData.CODE.value[index],
        header.get(JSONData.REPO.value, ''),
        set_raleway_medium, set_raleway
    )

    # Experience
    experiences = sections[JSONData.EXP.value]
    add_section(
        document, experiences[JSONData.TITLE.value],
        None, set_raleway_medium, set_raleway
    )
    create_experiences_table(
        document, experiences[JSONData.ITEMS.value],
        template_choice, index)

    # Education
    education = sections[JSONData.EDUCATION.value]
    add_section(
        document, f'\n{education[JSONData.TITLE.value]}',
        None, set_raleway_medium, set_raleway
    )
    add_bullet_list(
        document, education.get(JSONData.ITEMS.value, []),
        set_raleway, font_size=Style.NINE.value,
        bullet_color=Style.HEX_GREY.value
    )

    # Courses
    courses = sections[JSONData.COURSES.value]
    add_section(
        document, courses[JSONData.TITLE.value],
        None, set_raleway_medium, set_raleway
    )
    add_bullet_list(
        document, courses.get(JSONData.ITEMS.value, []),
        set_raleway, font_size=Style.NINE.value,
        bullet_color=Style.HEX_GREY.value
    )

    document.save(output_stream)


def _extract_table_text(table):
    """Retrieve text from table."""
    table_text = []
    for row in table.rows:
        row_text = [cell.text.strip() if cell.text.strip() else "-" for cell in row.cells]  # "-" для пустых ячеек
        if any(cell != "-" for cell in row_text):  # Добавляем только строки с данными
            table_text.append("\t".join(row_text))  # Объединяем ячейки через табуляцию
    return "\n".join(table_text).strip()  # Возвращаем строку с переносами строк между рядами


def extract_text_from_docx(doc_path):
    """
    DOCX parser: retrieve text, lists, tables.
    """
    if not os.path.exists(doc_path):
        logger.error(f"File {doc_path} not found.")
        return "Error: File not found."

    doc = Document(doc_path)
    extracted_text = []

    # Обработка параграфов
    for para in doc.paragraphs:
        if para.text.strip():  # Проверка, что текст не пустой
            extracted_text.append(para.text)

    # Обработка таблиц
    for table in doc.tables:
        table_text = _extract_table_text(table)
        if table_text.strip():  # Проверка на пустоту таблицы
            extracted_text.append(table_text)

    if not extracted_text:
        logger.error("File is empty.")
        return "Error: No readable content found."

    return "\n".join(extracted_text).strip()  # Объединяем все текстовые данные в одну строку

