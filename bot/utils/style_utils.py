from typing import Union

from docx.document import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, RGBColor
from docx.text.paragraph import Paragraph

from common.enums import Number, Style


def add_images_to_header_footer(
        document: Document, header_image_path: str,
        footer_image_path: str
) -> None:
    """
    Add header and footer images to the document.
    """

    section = document.sections[Number.ZERO.value]

    # Add header
    header = section.header
    header_paragraph = (
        header.paragraphs[
            Number.ZERO.value
        ] if header.paragraphs else header.add_paragraph()
    )
    header_paragraph.clear()
    header_paragraph.add_run().add_picture(
        header_image_path,
        width=Cm(Style.HEADER_WD.value),
        height=Cm(Style.HEADER_H.value)
    )

    # Add footer
    footer = section.footer
    footer_paragraph = (
        footer.paragraphs[
            Number.ZERO.value
        ] if footer.paragraphs else footer.add_paragraph()
    )
    footer_paragraph.clear()
    footer_paragraph.add_run().add_picture(
        footer_image_path,
        width=Cm(Style.FOOTER_WD.value),
        height=Cm(Style.FOOTER_H.value)
    )


def set_marker_style(
        paragraph: Paragraph, color: Union[RGBColor, str], font_size: int
) -> None:
    """
    Set the style for a paragraph marker with specified color and font size.
    """
    if isinstance(color, RGBColor):
        color = str(color)

    font_size_twips = str(int(font_size * Number.TWO.value))

    pPr = paragraph._element.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr')) or OxmlElement('w:numPr')
    pPr.append(numPr)

    rPr = pPr.find(qn('w:rPr')) or OxmlElement('w:rPr')
    pPr.append(rPr)

    color_element = rPr.find(qn('w:color')) or OxmlElement('w:color')
    rPr.append(color_element)
    color_element.set(qn('w:val'), color)

    size_element = rPr.find(qn('w:sz')) or OxmlElement('w:sz')
    rPr.append(size_element)
    size_element.set(qn('w:val'), font_size_twips)


def add_style_table(table, top_border):
    """
    Add style to table, making only the top border visible.
    """
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._element.get_or_add_tcPr()
            tc.append(parse_xml(r'''
                <w:tcBorders {}>
                    {border}
                    <w:left w:val="nil" w:sz="0" w:space="0" w:color="auto"/>
                    <w:bottom w:val="nil" w:sz="0" w:space="0" w:color="auto"/>
                    <w:right w:val="nil" w:sz="0" w:space="0" w:color="auto"/>
                </w:tcBorders>
            '''.format(nsdecls('w'), border=top_border)))
