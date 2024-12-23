from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from common.enums import JSONData, Style, Table
from .font_utils import set_raleway, set_raleway_medium
from .style_utils import set_marker_style


def create_skills_table(document, skills):
    """
    Creates table with skills, evenly distributed across columns
    (maximum 3 columns).
    """
    if not skills:
        return

    sorted_skills = sorted(skills)
    max_columns = 3
    rows = (len(sorted_skills) + max_columns - 1) // max_columns

    table = document.add_table(rows=rows, cols=max_columns)
    table.style = Table.STYLE.value

    # Заполняем таблицу
    skill_index = 0
    for row in table.rows:
        for cell in row.cells:
            if skill_index < len(sorted_skills):
                skill = sorted_skills[skill_index]
                add_skill_to_cell(cell, skill)
                skill_index += 1
            else:
                cell.text = ''

    style_table(table, Table.TRANSCEND.value)


def add_skill_to_cell(cell, skill):
    """
    Adds skill to the given table cell with style and centering.
    """
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    skill_run = paragraph.add_run(skill)
    set_raleway(skill_run, font_size=9)


def style_table(table, top_border):
    """
    Styles the table, making only the top border visible.
    """
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._element.get_or_add_tcPr()
            tc.append(parse_xml(r'''
                <w:tcBorders {}>
                    {border}
                    <w:left w:val="nil" w:sz="0" w:space="0" w:color="auto"/>
                    <w:bottom w:val="nil" w:sz="0" w:space="0" w:color="auto"/>
                    <w:right w:val="nil" w:sz="0" w:space="0" w:color="auto"/>
                </w:tcBorders>
            '''.format(nsdecls('w'), border=top_border)))


def create_experiences_table(document, experiences):
    """
    Creates table with experience with 2 columns.
    """
    # Create the table
    table = document.add_table(rows=0, cols=2)
    table.style = Table.STYLE.value

    # Iterate through experiences and populate the table
    for experience in experiences:
        row = table.add_row()
        left_cell, right_cell = row.cells

        # Left cell: Job title, period, company, and location
        left_paragraph = left_cell.paragraphs[0]

        role_data = left_paragraph.add_run(
            f'{experience.get(JSONData.ROLE.value, "")}\n'
        )
        set_raleway_medium(role_data, font_size=9)

        dates_data = left_paragraph.add_run(
            experience.get(JSONData.DATES.value)
        )
        set_raleway(dates_data, font_size=9)

        # Right cell: Project details
        right_paragraph = right_cell.paragraphs[0]

        # Add project title and description
        for title, data, new_line in JSONData.PROJECT.value:
            project_title = right_paragraph.add_run(title)
            set_raleway_medium(project_title, font_size=9)

            project_data = right_paragraph.add_run(experience.get(data, ''))
            set_raleway(project_data, font_size=9)

            if new_line:
                right_paragraph.add_run('\n')


        # Add tasks, achievements as bulleted lists
        for title, data in JSONData.TASKS_ACHIEVEMENTS.value:
            section_paragraph = right_cell.add_paragraph()
            section_title = section_paragraph.add_run(title)
            set_raleway_medium(section_title, font_size=9)

            for d in experience.get(data, []):
                bullet = right_cell.add_paragraph(style=Style.BULLET.value)
                set_marker_style(bullet, Style.HEX_GREY.value, 9)
                bullet_data = bullet.add_run(d)
                set_raleway(bullet_data, font_size=9)

        # Add team and stack
        for title, data in JSONData.TEAM_STACK.value:
            section_paragraph = right_cell.add_paragraph()
            section_title = section_paragraph.add_run(title)
            set_raleway_medium(section_title, font_size=9)

            section_data = section_paragraph.add_run(experience.get(data, ''))
            set_raleway(section_data, font_size=9)

    # Set table border styles (only top border visible)
    style_table(table, Table.TOP.value)
